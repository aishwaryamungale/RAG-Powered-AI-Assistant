import os
import glob
from PyPDF2 import PdfReader
from langchain_core.documents import Document


class DocumentProcessor:
    def __init__(self, data_dir="./data"):
        self.data_dir = data_dir
        os.makedirs(data_dir, exist_ok=True)
    
    def load_documents(self):
        documents = []
        documents.extend(self._load_pdfs())
        documents.extend(self._load_txts())
        documents.extend(self._load_docxs())
        return documents
    
    def _load_pdfs(self):
        pdf_files = glob.glob(os.path.join(self.data_dir, "*.pdf"))
        documents = []
        for pdf_file in pdf_files:
            try:
                reader = PdfReader(pdf_file)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": pdf_file}))
            except Exception as e:
                print(f"PDF read error {pdf_file}: {e}")
        return documents
    
    def _load_txts(self):
        txt_files = glob.glob(os.path.join(self.data_dir, "*.txt"))
        documents = []
        for txt_file in txt_files:
            try:
                with open(txt_file, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                    if text.strip():
                        documents.append(Document(page_content=text, metadata={"source": txt_file}))
            except Exception as e:
                print(f"TXT read error {txt_file}: {e}")
        return documents
    
    def _load_docxs(self):
        docx_files = glob.glob(os.path.join(self.data_dir, "*.docx"))
        documents = []
        try:
            from docx import Document as DocxDocument
            for docx_file in docx_files:
                try:
                    doc = DocxDocument(docx_file)
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    if text.strip():
                        documents.append(Document(page_content=text, metadata={"source": docx_file}))
                except Exception as e:
                    print(f"DOCX read error {docx_file}: {e}")
        except ImportError:
            print("python-docx not installed. DOCX files will be skipped.")
        return documents